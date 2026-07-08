#!/usr/bin/env python3
"""CyberStrikeAI docx 工具:读取模板结构 / 列模板 / 原地填充。
成功时向 stdout 输出一行 JSON;失败时向 stderr 写错误并 exit(1)。
"""
import argparse
import json
import os
import sys

from docx import Document

# 封面字段样式(style_id;真实模板中 name 为 "封面副标题（GYKJ）")
COVER_STYLE_ID = "GYKJf1"
# 章节提示语常用前缀
HINT_PREFIXES = ("编写说明", "提示")


def _open_document(path):
    """打开 docx 或 dotx 文件,兼容模板格式。"""
    try:
        return Document(path)
    except ValueError as e:
        if "template" not in str(e):
            raise
    import shutil
    import tempfile
    import zipfile
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp_path = tmp.name
    tmp.close()
    shutil.copy2(path, tmp_path)
    with zipfile.ZipFile(tmp_path, "r") as zin:
        ct = zin.read("[Content_Types].xml").decode("utf-8")
    ct_fixed = ct.replace(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
    )
    out_path = tmp_path + ".fixed.docx"
    with zipfile.ZipFile(tmp_path, "r") as zin:
        with zipfile.ZipFile(out_path, "w") as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = ct_fixed.encode("utf-8")
                zout.writestr(item, data)
    os.unlink(tmp_path)
    doc = Document(out_path)
    os.unlink(out_path)
    return doc


def _para_text(p):
    return (p.text or "").strip()


def _is_heading1(style):
    """判断是否为一级标题样式(兼容不同模板的 style_id)。"""
    if style is None:
        return False
    sid = getattr(style, "style_id", "") or ""
    name = getattr(style, "name", "") or ""
    if sid == "1GYKJ" or sid == "188":
        return True
    if "标题 1" in name or "Heading 1" in name.lower():
        return True
    return False


def _style_matches(style, target_id):
    if style is None:
        return False
    return style.style_id == target_id or style.name == target_id


def cmd_outline(template):
    doc = _open_document(template)
    headings = []
    tables_meta = []
    cover_fields = []

    paras = doc.paragraphs
    for idx, p in enumerate(paras):
        text = _para_text(p)
        if not text:
            continue
        if _is_heading1(p.style):
            hint = ""
            for nxt in paras[idx + 1:]:
                if _is_heading1(nxt.style):
                    break
                ntext = _para_text(nxt)
                if ntext.startswith(HINT_PREFIXES):
                    hint = ntext
                    break
            style_id = getattr(p.style, "style_id", "") or ""
            headings.append(
                {"anchor": text, "style": style_id, "text": text, "hint": hint}
            )
        elif _style_matches(p.style, COVER_STYLE_ID) and text.endswith("："):
            cover_fields.append(text.rstrip("："))

    for i, tbl in enumerate(doc.tables):
        columns = []
        if tbl.rows:
            columns = [c.text.strip() for c in tbl.rows[0].cells]
        name = columns[0] if columns else f"table_{i}"
        tables_meta.append({"index": i, "name": name, "columns": columns})

    return {"headings": headings, "tables": tables_meta, "cover_fields": cover_fields}


def cmd_list(root):
    templates = []
    for dirpath, _dirs, files in os.walk(root):
        for f in files:
            if (f.endswith(".docx") or f.endswith(".dotx")) and not f.startswith("~$"):
                full = os.path.join(dirpath, f)
                rel = os.path.relpath(full, root)
                templates.append({"path": rel, "name": os.path.splitext(f)[0]})
    templates.sort(key=lambda t: t["path"])
    return {"templates": templates}


def main():
    parser = argparse.ArgumentParser()
    sub = parser.add_subparsers(dest="cmd", required=True)

    p_out = sub.add_parser("outline")
    p_out.add_argument("--template", required=True)

    p_list = sub.add_parser("list")
    p_list.add_argument("--root", required=True)

    p_fill = sub.add_parser("fill")
    p_fill.add_argument("--payload", required=True)

    args = parser.parse_args()
    try:
        if args.cmd == "outline":
            result = cmd_outline(args.template)
        elif args.cmd == "list":
            result = cmd_list(args.root)
        elif args.cmd == "fill":
            from docx_fill import cmd_fill  # Task 2 提供
            result = cmd_fill(args.payload)
        else:
            raise SystemExit(2)
    except Exception as exc:  # noqa: BLE001
        sys.stderr.write(str(exc))
        sys.exit(1)
    sys.stdout.write(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
