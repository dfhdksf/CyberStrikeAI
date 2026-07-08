"""docx 原地填充:在保留模板全部结构的前提下,把内容插入指定章节与表格。

设计要点:
- 章节:定位 style_id=1GYKJ 的一级标题段落,在其后插入正文段落(style_id=GYKJ5)
  与项目符号段落,原有标题、封面、页眉页脚、目录、表格全部保留。
- Markdown 智能映射:paragraphs 中的 ## 映射为二级标题(2GYKJ),### 映射为三级标题(3GYKJ),
  - / * 映射为列表(GYKJ1),**bold** 映射为 run 级别加粗,普通文本为正文(GYKJ5)。
- 表格:按 index 填充——优先复用已有空行,不够再追加新行。
- 封面字段:真实模板中封面字段(文档编号 / 文档密级 / 当前版本 / 发布时间 等)
  以形如 "字段名：" 的形式存放在 table[0] 的单元格里,而不是独立段落。
  因此填充时必须同时扫描 doc.paragraphs 与所有 table.cells。
"""
import copy
import json
import re

from docx import Document
from docx.text.paragraph import Paragraph

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
HEADING_STYLE_ID = "1GYKJ"
HEADING2_STYLE_ID = "2GYKJ"
HEADING3_STYLE_ID = "3GYKJ"
BODY_STYLE_ID = "GYKJ5"
BULLET_STYLE_ID = "GYKJ1"
BULLET_STYLE = "List Bullet"


def _open_document(path):
    """打开 docx 或 dotx 文件,兼容模板格式。"""
    try:
        return Document(path)
    except ValueError as e:
        if "template" not in str(e):
            raise
    import shutil
    import tempfile
    import zipfile
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    tmp_path = tmp.name
    tmp.close()
    shutil.copy2(path, tmp_path)
    with zipfile.ZipFile(tmp_path, "r") as zin:
        ct = zin.read("[Content_Types].xml").decode("utf-8")
    ct_fixed = ct.replace(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
    )
    out_path = tmp_path + ".fixed.docx"
    with zipfile.ZipFile(tmp_path, "r") as zin:
        with zipfile.ZipFile(out_path, "w") as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = ct_fixed.encode("utf-8")
                zout.writestr(item, data)
    import os
    os.unlink(tmp_path)
    doc = Document(out_path)
    os.unlink(out_path)
    return doc


def _is_heading1(style):
    """判断是否为一级标题样式(兼容不同模板的 style_id)。"""
    if style is None:
        return False
    sid = getattr(style, "style_id", "") or ""
    name = getattr(style, "name", "") or ""
    if sid == "1GYKJ" or sid == "188":
        return True
    if "标题 1" in name or "Heading 1" in name.lower():
        return True
    return False


def _resolve_style(doc, style_key):
    """按 style_id 或 name 解析样式;找不到返回 None(调用方回退默认)。"""
    if not style_key:
        return None
    for style in doc.styles:
        if getattr(style, "style_id", None) == style_key:
            return style
    try:
        return doc.styles[style_key]
    except KeyError:
        return None


def _apply_style(para, doc, style_key):
    style = _resolve_style(doc, style_key)
    if style is not None:
        try:
            para.style = style
        except (KeyError, ValueError):
            pass


def _insert_paragraph_after(paragraph, text, doc, style_key):
    """在给定段落之后插入一个新段落(克隆 pPr,保持文档流与段落属性)。"""
    new_p = copy.deepcopy(paragraph._p)
    for r in new_p.findall(W_NS + "r"):
        new_p.remove(r)
    paragraph._p.addnext(new_p)

    para = Paragraph(new_p, paragraph._parent)
    _apply_style(para, doc, style_key)
    para.add_run(text)
    return para


def _add_runs_with_inline(para, text):
    """解析 inline markdown(**bold**),拆分为多个 run。"""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part:
            para.add_run(part)


def _resolve_heading_style(doc, level):
    """解析指定层级的标题样式,兼容不同模板的 style_id 命名。"""
    candidates_by_level = {
        2: ["2GYKJ", "189"],
        3: ["3GYKJ", "190"],
    }
    for sid in candidates_by_level.get(level, []):
        style = _resolve_style(doc, sid)
        if style:
            return sid
    # 按名称回退
    name_pattern = f"标题 {level}"
    for style in doc.styles:
        if name_pattern in (getattr(style, "name", "") or ""):
            return getattr(style, "style_id", None)
    return BODY_STYLE_ID


def _resolve_bullet_style(doc):
    """解析列表样式,兼容不同模板。"""
    for sid in ["GYKJ1", "GYKJ"]:
        if _resolve_style(doc, sid):
            return sid
    if _resolve_style(doc, BULLET_STYLE):
        return BULLET_STYLE
    return BODY_STYLE_ID


def _parse_and_insert(anchor_p, text, doc):
    """解析单行 markdown 标记,映射为对应 docx 样式插入。"""
    # ## 二级标题
    m = re.match(r"^##\s+(.+)$", text)
    if m:
        h2_style = _resolve_heading_style(doc, 2)
        return _insert_paragraph_after(anchor_p, m.group(1), doc, h2_style)

    # ### 三级标题
    m = re.match(r"^###\s+(.+)$", text)
    if m:
        h3_style = _resolve_heading_style(doc, 3)
        return _insert_paragraph_after(anchor_p, m.group(1), doc, h3_style)

    # - / * / + 无序列表
    m = re.match(r"^[-*+]\s+(.+)$", text)
    if m:
        content = m.group(1)
        bullet_style = _resolve_bullet_style(doc)
        para = _insert_paragraph_after(anchor_p, "", doc, bullet_style)
        para.runs[0].text = ""
        _add_runs_with_inline(para, content)
        return para

    # 数字编号列表 (1. / 2. / ...)
    m = re.match(r"^\d+\.\s+(.+)$", text)
    if m:
        content = m.group(1)
        bullet_style = _resolve_bullet_style(doc)
        para = _insert_paragraph_after(anchor_p, "", doc, bullet_style)
        para.runs[0].text = ""
        _add_runs_with_inline(para, content)
        return para

    # 普通正文段落(含 inline bold 解析)
    body_style = _resolve_style(doc, BODY_STYLE_ID) and BODY_STYLE_ID or "Normal"
    para = _insert_paragraph_after(anchor_p, "", doc, body_style)
    para.runs[0].text = ""
    _add_runs_with_inline(para, text)
    return para


def _iter_all_paragraphs(doc):
    """遍历文档 body 段落与全部表格单元格内段落。"""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _find_heading(doc, anchor):
    for p in doc.paragraphs:
        if _is_heading1(p.style) and anchor in (p.text or ""):
            return p
    return None


def _label_of(text):
    """把段落/单元格文本规范化为封面字段标签(去空白 + 去掉尾部中英文冒号)。"""
    return (text or "").strip().rstrip("：:").strip()


def _fill_cover(doc, cover):
    """把封面字段值追加到匹配的段落 / 表格单元格。"""
    if not cover:
        return
    # 同一 payload 中同名字段只填一次,避免误伤重复出现的标签
    remaining = dict(cover)

    for p in _iter_all_paragraphs(doc):
        if not remaining:
            return
        label = _label_of(p.text)
        if label and label in remaining:
            p.add_run(" " + str(remaining.pop(label)))


def cmd_fill(payload_path):
    with open(payload_path, "r", encoding="utf-8") as f:
        payload = json.load(f)

    doc = _open_document(payload["template"])
    filled_sections = 0
    filled_tables = 0

    # 章节填充:定位标题后按序插入(自动解析 markdown 映射 docx 样式)
    for sec in (payload.get("sections") or []):
        heading = _find_heading(doc, sec["anchor"])
        if heading is None:
            continue
        anchor_p = heading
        for text in sec.get("paragraphs", []):
            anchor_p = _parse_and_insert(anchor_p, text, doc)
        bullet_style = _resolve_bullet_style(doc)
        for bullet in sec.get("bullets", []):
            anchor_p = _insert_paragraph_after(anchor_p, bullet, doc, bullet_style)
        filled_sections += 1

    # 表格填充:按 index 填充——优先复用已有空行,不够再追加
    for t in (payload.get("tables") or []):
        idx = t.get("index", -1)
        if idx < 0 or idx >= len(doc.tables):
            continue
        table = doc.tables[idx]
        ncols = len(table.columns)
        fill_row = 1  # 跳过 header row[0]
        for row_vals in t.get("rows", []):
            if fill_row < len(table.rows):
                cells = table.rows[fill_row].cells
            else:
                cells = table.add_row().cells
            for ci, val in enumerate(row_vals[:ncols]):
                cells[ci].text = str(val)
            fill_row += 1
        filled_tables += 1

    # 封面填充(段落 + 表格单元格两种存放形式都要处理)
    _fill_cover(doc, payload.get("cover") or {})

    doc.save(payload["output"])
    return {
        "output_path": payload["output"],
        "filled_sections": filled_sections,
        "filled_tables": filled_tables,
    }
