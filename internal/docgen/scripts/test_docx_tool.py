import json
import os
import subprocess
import unittest

HERE = os.path.dirname(os.path.abspath(__file__))
SCRIPT = os.path.join(HERE, "docx_tool.py")
REPO_ROOT = os.path.abspath(os.path.join(HERE, "..", "..", ".."))
TEMPLATE_ROOT = os.path.join(REPO_ROOT, "reference", "表单模板")
TECH_TEMPLATE = os.path.join(
    TEMPLATE_ROOT, "1.项目立项", "GYAQ-QM-IPM-MT-02.技术预研计划.docx"
)


def run(*args):
    proc = subprocess.run(
        ["python3", SCRIPT, *args], capture_output=True, text=True
    )
    return proc


class TestOutline(unittest.TestCase):
    def test_outline_returns_five_headings(self):
        proc = run("outline", "--template", TECH_TEMPLATE)
        self.assertEqual(proc.returncode, 0, proc.stderr)
        data = json.loads(proc.stdout)
        texts = [h["text"] for h in data["headings"]]
        # 5 个正文一级章节(样式 1GYKJ)
        self.assertIn("技术预研目标", "".join(texts))
        self.assertIn("可能存在的困难与风险", "".join(texts))
        self.assertGreaterEqual(len([h for h in data["headings"] if h["style"] == "1GYKJ"]), 5)

    def test_outline_captures_hint_text(self):
        proc = run("outline", "--template", TECH_TEMPLATE)
        data = json.loads(proc.stdout)
        joined = json.dumps(data, ensure_ascii=False)
        # 章节下的"编写说明/提示"文本被采集为 hint
        self.assertIn("可以验证", joined)

    def test_outline_lists_tables(self):
        proc = run("outline", "--template", TECH_TEMPLATE)
        data = json.loads(proc.stdout)
        self.assertGreaterEqual(len(data["tables"]), 5)


class TestList(unittest.TestCase):
    def test_list_finds_templates(self):
        proc = run("list", "--root", TEMPLATE_ROOT)
        self.assertEqual(proc.returncode, 0, proc.stderr)
        data = json.loads(proc.stdout)
        paths = [t["path"] for t in data["templates"]]
        self.assertTrue(any("技术预研计划" in p for p in paths))
        self.assertTrue(all(p.endswith(".docx") or p.endswith(".dotx") for p in paths))


class TestFill(unittest.TestCase):
    def _fill_payload(self, out_path):
        return {
            "template": TECH_TEMPLATE,
            "output": out_path,
            "sections": [
                {
                    "anchor": "技术预研目标",
                    "paragraphs": ["验证 docx 原地填充可行性。"],
                    "bullets": ["目标一", "目标二"],
                }
            ],
            "tables": [],
            "cover": {"当前版本": "V1.0"},
        }

    def test_fill_creates_output_and_inserts_content(self):
        import tempfile
        from docx import Document

        with tempfile.TemporaryDirectory() as td:
            out = os.path.join(td, "out.docx")
            payload = os.path.join(td, "p.json")
            with open(payload, "w", encoding="utf-8") as f:
                json.dump(self._fill_payload(out), f, ensure_ascii=False)
            proc = run("fill", "--payload", payload)
            self.assertEqual(proc.returncode, 0, proc.stderr)
            res = json.loads(proc.stdout)
            self.assertTrue(os.path.exists(res["output_path"]))
            # 内容被插入
            doc = Document(out)
            body = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("验证 docx 原地填充可行性", body)
            self.assertIn("目标一", body)

    def test_fill_preserves_headers_and_tables(self):
        import tempfile
        from docx import Document

        original = Document(TECH_TEMPLATE)
        orig_table_count = len(original.tables)
        orig_section_count = len(original.sections)

        with tempfile.TemporaryDirectory() as td:
            out = os.path.join(td, "out.docx")
            payload = os.path.join(td, "p.json")
            with open(payload, "w", encoding="utf-8") as f:
                json.dump(self._fill_payload(out), f, ensure_ascii=False)
            run("fill", "--payload", payload)
            doc = Document(out)
            # 表格数量不变(封面/版本表/进度表等全部保留)
            self.assertEqual(len(doc.tables), orig_table_count)
            # 分节数量不变(封面/目录/正文分节保留 → 页眉页脚随之保留)
            self.assertEqual(len(doc.sections), orig_section_count)

    def test_fill_cover_field_in_table_gets_written(self):
        """封面字段位于 table[0] 单元格中,填充后单元格文本应包含值。"""
        import tempfile
        from docx import Document

        with tempfile.TemporaryDirectory() as td:
            out = os.path.join(td, "out.docx")
            payload_data = self._fill_payload(out)
            # 只保留 cover 字段测试,清空 sections 干扰
            payload_data["cover"] = {"当前版本": "V1.0"}
            payload = os.path.join(td, "p.json")
            with open(payload, "w", encoding="utf-8") as f:
                json.dump(payload_data, f, ensure_ascii=False)
            proc = run("fill", "--payload", payload)
            self.assertEqual(proc.returncode, 0, proc.stderr)
            doc = Document(out)
            # 值必须真的落到表格 0 里,证明我们处理了 cell 中的封面字段
            table_text = "\n".join(
                cell.text
                for row in doc.tables[0].rows
                for cell in row.cells
            )
            self.assertIn("V1.0", table_text)


class TestFillMarkdownMapping(unittest.TestCase):

    def _run_fill_with_paragraphs(self, paragraphs):
        import tempfile
        from docx import Document

        with tempfile.TemporaryDirectory() as td:
            out = os.path.join(td, "out.docx")
            payload_data = {
                "template": TECH_TEMPLATE,
                "output": out,
                "sections": [
                    {
                        "anchor": "技术预研目标",
                        "paragraphs": paragraphs,
                        "bullets": [],
                    }
                ],
                "tables": [],
                "cover": {},
            }
            payload = os.path.join(td, "p.json")
            with open(payload, "w", encoding="utf-8") as f:
                json.dump(payload_data, f, ensure_ascii=False)
            proc = run("fill", "--payload", payload)
            self.assertEqual(proc.returncode, 0, proc.stderr)
            return Document(out)

    def test_h2_mapped_to_heading2_style(self):
        doc = self._run_fill_with_paragraphs(["## 研究背景"])
        found = False
        for p in doc.paragraphs:
            if "研究背景" in (p.text or ""):
                sid = getattr(p.style, "style_id", "")
                if sid == "2GYKJ":
                    found = True
                    break
        self.assertTrue(found, "## should map to 2GYKJ style")

    def test_h3_mapped_to_heading3_style(self):
        doc = self._run_fill_with_paragraphs(["### 细分方向"])
        found = False
        for p in doc.paragraphs:
            if "细分方向" in (p.text or ""):
                sid = getattr(p.style, "style_id", "")
                if sid == "3GYKJ":
                    found = True
                    break
        self.assertTrue(found, "### should map to 3GYKJ style")

    def test_bullet_mapped_to_list_style(self):
        doc = self._run_fill_with_paragraphs(["- 要点一", "* 要点二"])
        bullets_found = 0
        for p in doc.paragraphs:
            if p.text in ("要点一", "要点二"):
                sid = getattr(p.style, "style_id", "")
                if sid == "GYKJ1":
                    bullets_found += 1
        self.assertEqual(bullets_found, 2, "- and * should map to GYKJ1 list style")

    def test_bold_inline_creates_bold_run(self):
        doc = self._run_fill_with_paragraphs(["这是**重点**内容"])
        for p in doc.paragraphs:
            if "重点" in (p.text or ""):
                for run in p.runs:
                    if run.text == "重点":
                        self.assertTrue(run.bold, "**text** should produce bold run")
                        return
        self.fail("Bold run not found")

    def test_no_markdown_symbols_in_output_text(self):
        doc = self._run_fill_with_paragraphs([
            "## 子标题", "- 列表项", "**加粗**文本"
        ])
        body = "\n".join(p.text for p in doc.paragraphs)
        self.assertNotIn("##", body)
        self.assertNotIn("- 列表项", body)
        self.assertNotIn("**", body)
        self.assertIn("子标题", body)
        self.assertIn("列表项", body)
        self.assertIn("加粗", body)


class TestFillTablePosition(unittest.TestCase):
    """验证表格填充复用已有空行而不是追加到末尾。"""

    def test_table_fills_existing_empty_rows(self):
        import tempfile
        from docx import Document

        with tempfile.TemporaryDirectory() as td:
            out = os.path.join(td, "out.docx")
            payload_data = {
                "template": TECH_TEMPLATE,
                "output": out,
                "sections": [],
                "tables": [
                    {
                        "index": 4,
                        "rows": [
                            ["任务A", "2026-01-01", "2026-02-01", "张三"],
                            ["任务B", "2026-02-01", "2026-03-01", "李四"],
                        ],
                    }
                ],
                "cover": {},
            }
            payload = os.path.join(td, "p.json")
            with open(payload, "w", encoding="utf-8") as f:
                json.dump(payload_data, f, ensure_ascii=False)
            proc = run("fill", "--payload", payload)
            self.assertEqual(proc.returncode, 0, proc.stderr)
            doc = Document(out)
            table = doc.tables[4]
            # row[0] is header, data should be in row[1] and row[2]
            self.assertIn("任务A", table.rows[1].cells[0].text)
            self.assertIn("任务B", table.rows[2].cells[0].text)

    def test_table_overflow_appends_new_rows(self):
        import tempfile
        from docx import Document

        orig = Document(TECH_TEMPLATE)
        orig_row_count = len(orig.tables[4].rows)

        with tempfile.TemporaryDirectory() as td:
            out = os.path.join(td, "out.docx")
            # 填充超过空行数量的数据(模板有3个空行,填5行数据)
            rows = [[f"任务{i}", f"开始{i}", f"结束{i}", f"人员{i}"] for i in range(5)]
            payload_data = {
                "template": TECH_TEMPLATE,
                "output": out,
                "sections": [],
                "tables": [{"index": 4, "rows": rows}],
                "cover": {},
            }
            payload = os.path.join(td, "p.json")
            with open(payload, "w", encoding="utf-8") as f:
                json.dump(payload_data, f, ensure_ascii=False)
            proc = run("fill", "--payload", payload)
            self.assertEqual(proc.returncode, 0, proc.stderr)
            doc = Document(out)
            table = doc.tables[4]
            # 原来4行(1 header + 3 empty), 填5行数据 → 应该6行(1 header + 5 data)
            self.assertEqual(len(table.rows), orig_row_count + 2)
            # 最后一行应该有数据
            self.assertIn("任务4", table.rows[5].cells[0].text)


if __name__ == "__main__":
    unittest.main()

